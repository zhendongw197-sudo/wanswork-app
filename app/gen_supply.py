# -*- coding: utf-8 -*-
"""供方单（电力用户新装供电方案确认单）生成：红字占位模板替换。

模板（assets/供电方案确认单.docx，由 tools/build_supply_template.py 重建）中
需替换的文字为红色（FF0000）独立 run；本模块遍历正文段落及所有表格
（含嵌套）单元格中的 run，颜色为红的按 app/supply.py 的 build_field_map
查表替换文字并改黑；未命中的占位名保留原文并记入返回列表。

依赖：python-docx。
"""
from __future__ import annotations

import os

from docx import Document
from docx.shared import RGBColor

try:
    from .supply import build_field_map as _build_field_map
except ImportError:  # app/supply.py 由另一成员并行开发，尚未就绪
    _build_field_map = None

_BLACK = RGBColor(0x00, 0x00, 0x00)

# 纯标点占位（如模板中误标红的顿号'、'）：替换时直接改黑，不报警
_PUNCT_ONLY = set('、，。；：！？,.;:!?（）()【】[] ')


def _is_red(run) -> bool:
    """判断 run 是否为红色占位（color.rgb 可能为 None 或主题色，需容错）。"""
    try:
        rgb = run.font.color.rgb
    except Exception:
        return False
    return rgb is not None and str(rgb).upper() == 'FF0000'


def _iter_table_paragraphs(table):
    """递归产出表格（含嵌套表）每个单元格中的段落。"""
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                yield p
            for t in cell.tables:
                yield from _iter_table_paragraphs(t)


def _iter_paragraphs(doc):
    """产出文档正文段落 + 所有表格单元格段落。"""
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        yield from _iter_table_paragraphs(table)


def gen_supply_form(order: dict, supply: dict, template_path: str,
                    out_path: str, field_map: dict | None = None) -> list[str]:
    """按模板生成供方单，返回未匹配的红字占位名列表（空=全部替换成功）。

    field_map 仅供测试注入；默认取 app/supply.py 的 build_field_map。
    """
    if field_map is None:
        if _build_field_map is None:
            raise ImportError('app/supply.py 尚未就绪，且未传入 field_map')
        field_map = _build_field_map(order, supply)

    doc = Document(template_path)
    unmatched: list[str] = []
    for p in _iter_paragraphs(doc):
        for run in p.runs:
            if not _is_red(run):
                continue
            name = run.text.strip()
            if name in field_map:
                run.text = str(field_map[name])
                run.font.color.rgb = _BLACK
            elif name and all(ch in _PUNCT_ONLY for ch in name):
                # 纯标点误标红：保留文字、改黑、不记入未匹配列表
                run.font.color.rgb = _BLACK
            else:
                unmatched.append(name)

    os.makedirs(os.path.dirname(os.path.abspath(out_path)), exist_ok=True)
    doc.save(out_path)
    return unmatched
