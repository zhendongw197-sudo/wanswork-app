# -*- coding: utf-8 -*-
"""营销工单配套 Word 材料生成（情况说明 / 交底申请书 / 施工方案 / 情况说明）。

依赖：python-docx。
order dict 约定键：flow_id, flow_type, account_no, account_name, start_time,
                  manager（姓名）, manager_phone, address, phone
"""
from __future__ import annotations

import os
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt

# 字号：小二 = 18pt，小四 = 12pt
_SIZE_TITLE = Pt(18)
_SIZE_BODY = Pt(12)


# ---------------------------------------------------------------------------
# 排版辅助
# ---------------------------------------------------------------------------
def _set_run_font(run, name: str, size, bold: bool = False):
    """同时设置西文与中文字体（python-docx 中文需单独设 eastAsia）。"""
    run.font.name = name
    run.font.size = size
    run.font.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), name)


def _add_title(doc: Document, text: str):
    """标题：居中、加粗、小二。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    _set_run_font(run, "宋体", _SIZE_TITLE, bold=True)
    return p


def _add_body(doc: Document, text: str, indent: bool = True, bold: bool = False):
    """正文：宋体小四，默认首行缩进 2 字符（24pt = 2 x 小四）。"""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    _set_run_font(run, "宋体", _SIZE_BODY, bold=bold)
    return p


def _add_signature(doc: Document, lines: list[str]):
    """落款：右对齐。"""
    for line in lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(line)
        _set_run_font(run, "宋体", _SIZE_BODY)


def _fmt_date_cn(value) -> str:
    """把日期值格式化为 XXXX年XX月XX日。支持 datetime 或 'YYYY-MM-DD...' 字符串。"""
    if isinstance(value, datetime):
        d = value
    else:
        d = datetime.strptime(str(value).strip()[:10], "%Y-%m-%d")
    return f"{d.year}年{d.month:02d}月{d.day:02d}日"


def _today_cn() -> str:
    return _fmt_date_cn(datetime.now())


def _save(doc: Document, out_path: str) -> str:
    os.makedirs(os.path.dirname(os.path.abspath(out_path)), exist_ok=True)
    doc.save(out_path)
    return out_path


# ---------------------------------------------------------------------------
# 材料4：关于"XX"工单申请终止白名单的情况说明
# ---------------------------------------------------------------------------
def gen_termination_note(order: dict, reason: str, out_path: str) -> str:
    doc = Document()
    _add_title(doc, f"关于“{order['account_name']}”工单申请终止白名单的情况说明")
    doc.add_paragraph()
    _add_body(doc, "一、流程信息", indent=False, bold=True)
    apply_date = _fmt_date_cn(order["start_time"])
    _add_body(doc, (
        f"户名：{order['account_name']}（户号：{order['account_no']}）于{apply_date}"
        f"申请{order['flow_type']}新装增容业务（流程编号：{order['flow_id']}），"
        f"用地地址：{order['address']}。"
    ))
    _add_body(doc, "二、原因描述", indent=False, bold=True)
    _add_body(doc, f"因用户{reason}，用户申请终止此流程。")
    doc.add_paragraph()
    _add_signature(doc, [_today_cn()])
    return _save(doc, out_path)


# ---------------------------------------------------------------------------
# 附件5：交底申请书
# ---------------------------------------------------------------------------
def gen_jiaodi_apply(order: dict, street: str, distance_m: str, jiaodi_date: str,
                     out_path: str, config: dict) -> str:
    doc = Document()
    _add_title(doc, "交底申请书")
    doc.add_paragraph()
    _add_body(doc, f"致 {config['gas_company']}：", indent=False)
    _add_body(doc, (
        f"位于{street}（街道）{order['account_name']}（项目）需要进行开挖施工，"
        f"施工范围在{order['address']}，需你公司配合交底，请于{jiaodi_date}在"
        f"{order['address']}携带地下燃气管线图等交底资料进行现场交底。"
    ))
    doc.add_paragraph()
    _add_signature(doc, [f"{config['org_name']}（单位）", _today_cn()])
    return _save(doc, out_path)


# ---------------------------------------------------------------------------
# 附件6：施工方案
# ---------------------------------------------------------------------------
def gen_construct_plan(order: dict, distance_m: str, out_path: str, config: dict) -> str:
    doc = Document()
    _add_title(doc, "二.施工方案")
    doc.add_paragraph()
    _add_body(doc, f"项目名称：{order['account_name']}业扩。", indent=False)
    _add_body(doc, (
        f"因{order['account_name']}用电新装，现需开挖敷设电缆。"
        f"开挖长度为{distance_m}米，开挖深度为{config['dig_depth_cm']}厘米。"
        f"施工范围：{order['address']}。该项目预计工期{config.get('period_days', '3')}天，"
        f"路面由供电公司负责恢复。"
    ))
    doc.add_paragraph()
    _add_body(doc, (
        f"建设单位：{config['org_name']}  负责人：{order['manager']}  "
        f"电话：{order['manager_phone']}"
    ), indent=False)
    _add_body(doc, (
        f"施工单位：{config['builder_name']}  负责人：{config['builder_leader']}  "
        f"电话：{config['builder_phone']}"
    ), indent=False)
    return _save(doc, out_path)


# ---------------------------------------------------------------------------
# 附件7：情况说明（无监理）
# ---------------------------------------------------------------------------
def gen_supervision_note(order: dict, out_path: str) -> str:
    doc = Document()
    _add_title(doc, "情况说明")
    doc.add_paragraph()
    _add_body(doc, (
        f"工程名称：{order['account_name']}（{order['address']}）{order['flow_type']}。"
    ))
    _add_body(doc, (
        "工程概况：此工程体量小技术简单，建设面积小，未设立监理，现场施工无监理。"
    ))
    doc.add_paragraph()
    _add_signature(doc, ["建设方", _today_cn()])
    return _save(doc, out_path)
