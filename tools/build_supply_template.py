# -*- coding: utf-8 -*-
"""重建供方单模板 assets/供电方案确认单.docx（电力用户新装供电方案确认单）。

原版为老 .doc，布局丢失，本脚本用 python-docx 按内容从零重建一页式表单。
占位约定：需替换的文字为红色（FF0000）独立 run，生成时由
app/gen_supply.py 按 build_field_map 查表替换并改黑。
"""
from __future__ import annotations

import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

_RED = RGBColor(0xFF, 0x00, 0x00)
_BLACK = RGBColor(0x00, 0x00, 0x00)

_SIZE_TITLE = Pt(22)   # 二号
_SIZE_BODY = Pt(12)    # 小四

_BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
TEMPLATE_PATH = os.path.join(_BASE, 'assets', '供电方案确认单.docx')


def _set_run_font(run, size, bold: bool = False, red: bool = False):
    """中文宋体（需单独设 eastAsia），占位 run 红色、其余黑色。"""
    run.font.name = '宋体'
    run.font.size = size
    run.font.bold = bold
    run.font.color.rgb = _RED if red else _BLACK
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = rpr.makeelement(qn('w:rFonts'), {})
        rpr.append(rfonts)
    rfonts.set(qn('w:eastAsia'), '宋体')


def _add_run(p, text: str, red: bool = False, bold: bool = False):
    """向段落追加一个宋体小四 run；red=True 表示红色占位。"""
    run = p.add_run(text)
    _set_run_font(run, _SIZE_BODY, bold=bold, red=red)
    return run


def _cell_write(cell, segments, bold: bool = False):
    """向单元格首段写入若干 (text, red) 片段。"""
    p = cell.paragraphs[0]
    for text, red in segments:
        _add_run(p, text, red=red, bold=bold)


def _set_borders(table):
    """给表格加全边框（不依赖内置样式，直接写 tblBorders）。"""
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.makeelement(qn('w:tblBorders'), {})
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = borders.makeelement(qn(f'w:{edge}'), {})
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:color'), '000000')
        borders.append(el)
    tbl_pr.append(borders)


def build_template(out_path: str) -> str:
    """从零重建模板并保存到 out_path，返回该路径。"""
    doc = Document()

    # 标题：居中、加粗、二号
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('电力用户新装供电方案确认单')
    _set_run_font(run, _SIZE_TITLE, bold=True)

    # 表格1（2行2列，无边框）：申请编号 / 申请类型
    t1 = doc.add_table(rows=2, cols=2)
    _cell_write(t1.cell(0, 0), [('申请编号：', False), ('流程编号', True)])
    _cell_write(t1.cell(0, 1), [('申请类型：', False), ('工单类型', True)])

    # 表格2（信息栏，带边框，4列：标签/值/标签/值）
    t2 = doc.add_table(rows=5, cols=4)
    _set_borders(t2)
    rows2 = [
        ([('户名', False)], [('户名', True)],
         [('户号', False)], [('户号', True)]),
        ([('行政区', False)], [('溧水区', False)],
         [('用电期限', False)], [('正式用电', False)]),
        ([('用电地址', False)], [('地址', True)],
         [('邮编', False)], [('211200', False)]),
        ([('联系人', False)], [('户名', True)],
         [('电话', False)], [('用户电话', True)]),
        ([('行业分类', False)], [('行业分类', True)],
         [('', False)], [('', False)]),
    ]
    for r, (c0, c1, c2, c3) in enumerate(rows2):
        for c, segs in enumerate((c0, c1, c2, c3)):
            _cell_write(t2.cell(r, c), segs)

    # 表格3（用检户信息，带边框，4列）
    t3 = doc.add_table(rows=6, cols=4)
    _set_borders(t3)
    rows3 = [
        ([('标志', False)], [('保留', False)],
         [('运行方式', False)], [('单电源', False)]),
        ([('电源性质', False)], [('主供电源', False)],
         [('线路变更情况', False)], [('', False)]),
        ([('供电电压', False)], [('供电电压', True), ('V', False)],
         [('批准容量', False)], [('批准容量', True), ('kVA', False)]),
        ([('供电线路', False)], [('线路', True)],
         [('公配名称', False)], [('变压器名称', True)]),
        ([('公配容量', False)], [('和同容量', True), ('kVA', False)],
         [('上一年最大负荷', False)], [('去年负荷', True), ('kW', False)]),
        ([('接户方式', False)], [('接户方式', True)],
         [('', False)], [('', False)]),
    ]
    for r, (c0, c1, c2, c3) in enumerate(rows3):
        for c, segs in enumerate((c0, c1, c2, c3)):
            _cell_write(t3.cell(r, c), segs)

    # 勘察员段落
    p = doc.add_paragraph()
    _add_run(p, '勘察员： ')
    _add_run(p, '户名', red=True)
    _add_run(p, '，申请')
    _add_run(p, '工单类型', red=True)
    _add_run(p, '，用于')
    _add_run(p, '行业分类', red=True)
    _add_run(p, '，合同容量为')
    _add_run(p, '批准容量', red=True)
    _add_run(p, 'kVA。')

    # 供电方案段落（'线路 变压器名称'为一个整体红 run）
    p = doc.add_paragraph()
    _add_run(p, '该户由')
    _add_run(p, '线路 变压器名称', red=True)
    _add_run(p, '供电，新放')
    _add_run(p, '施工规格', red=True)
    _add_run(p, '。新装')
    _add_run(p, '表箱型号', red=True)
    _add_run(p, '。计量装置配置')
    _add_run(p, '电表电压型号', red=True)
    _add_run(p, 'V伏、')
    _add_run(p, '电表电流型号', red=True)
    _add_run(p, '电表一只。执行')
    _add_run(p, '电价', red=True)
    _add_run(p, '。计量装置按照DL448规程和南京供电公司智能表配置要求设置，'
                '电力设施运行维护管理责任分界点设在计量箱出线开关下桩头以下10cm处。')

    # 表格4（意见栏，带边框，4行2列：左侧标签右侧留白）
    t4 = doc.add_table(rows=4, cols=2)
    _set_borders(t4)
    for r, label in enumerate(('低压客户经理班意见', '营销部意见',
                               '配电运检班意见', '运维检修部意见')):
        _cell_write(t4.cell(r, 0), [(label, False)])

    # 末段：右对齐日期
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _add_run(p, '年     月     日')

    os.makedirs(os.path.dirname(os.path.abspath(out_path)), exist_ok=True)
    doc.save(out_path)
    return out_path


if __name__ == '__main__':
    print(build_template(TEMPLATE_PATH))
