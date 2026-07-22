# -*- coding: utf-8 -*-
"""gen_supply 供方单生成测试：模板重建 + 红字占位替换。

app/supply.py 由另一成员并行开发；若尚未就绪，测试内联一份等价的
field_map 注入 gen_supply_form；就绪后自动改走契约路径 build_field_map。
"""

import os
import shutil
import subprocess
import sys
import tempfile

# 保证可以 import app 包（项目根目录加入 sys.path）
_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, _ROOT)

from docx import Document

from app.gen_supply import gen_supply_form

TEMPLATE_PATH = os.path.join(_ROOT, 'assets', '供电方案确认单.docx')

try:
    from app.supply import build_field_map as _contract_field_map
except ImportError:
    _contract_field_map = None


# ---------------------------------------------------------------------------
# 内联等价 field_map（契约见任务书；app/supply.py 就绪后不再使用）
# ---------------------------------------------------------------------------
def _meter_voltage(voltage: str) -> str:
    return '220' if voltage == '220' else '3×220/380'


def _compose_box_model(meter_current, box_phase, box_positions, box_mount) -> str:
    return f'{meter_current}{box_phase}{box_positions}{box_mount}表箱'


def _compose_materials_text(materials: dict) -> str:
    headers = {'cable_4x35': '4*35电缆', 'mpp100': 'MPP100管子',
               'term_35': '35冷缩终端'}
    parts = []
    for key, header in headers.items():
        v = materials.get(key, '')
        if v != '':
            unit = '只' if key.startswith('term_') else '米'
            parts.append(f'{header}{v}{unit}')
    return '、'.join(parts)


def _inline_field_map(order: dict, supply: dict) -> dict:
    return {
        '流程编号': order['flow_id'],
        '工单类型': order['flow_type'],
        '户名': order['account_name'],
        '户号': order['account_no'],
        '地址': order['address'],
        '用户电话': order['phone'],
        '行业分类': supply['industry'],
        '供电电压': supply['voltage'],
        '批准容量': supply['capacity'],
        '线路': supply['line'],
        '变压器名称': supply['transformer'],
        '和同容量': supply['contract_capacity'],
        '去年负荷': supply['last_year_load'],
        '接户方式': supply['connection'],
        '线路 变压器名称': f"{supply['line']} {supply['transformer']}",
        '施工规格': _compose_materials_text(supply['materials']),
        '表箱型号': _compose_box_model(supply['meter_current'], supply['box_phase'],
                                       supply['box_positions'], supply['box_mount']),
        '电表电压型号': _meter_voltage(supply['voltage']),
        '电表电流型号': supply['meter_current'],
        '电价': supply['price'],
    }


def _build_field_map(order: dict, supply: dict) -> dict | None:
    """优先走契约 build_field_map；未就绪时返回 None 表示需注入内联实现。"""
    if _contract_field_map is not None:
        return _contract_field_map(order, supply)
    return None


# ---------------------------------------------------------------------------
# 测试数据
# ---------------------------------------------------------------------------
def _fake_order() -> dict:
    return {
        'flow_id': 'GD20240108001',
        'flow_type': '低压非居民新装',
        'account_no': '3201240001234567',
        'account_name': '张三五金店',
        'address': '溧水区永阳街道xx路1号',
        'phone': '13800001111',
    }


def _fake_supply() -> dict:
    """380V、400A 电表、3 种材料的供方单业务字段。"""
    materials = {k: '' for k in (
        'cable_2x16', 'cable_4x16', 'cable_4x35', 'cable_4x70', 'cable_4x150',
        'cable_4x240', 'box_wall', 'box_floor', 'mpp100', 'tray_200x100',
        'tray_300x150', 'term_16', 'term_35', 'term_70', 'term_150',
        'term_240', 'overhead')}
    materials.update({'cable_4x35': 50, 'mpp100': 20, 'term_35': 2})
    return {
        'voltage': '380',
        'capacity': '200',
        'line': '10kV永阳线',
        'transformer': '永阳变2号公配',
        'contract_capacity': '200',
        'last_year_load': '120',
        'connection': '电缆',
        'meter_current': '400A',
        'price': '一般工商业电价',
        'industry': '商业',
        'box_phase': '三相',
        'box_positions': '一表位',
        'box_mount': '落地式',
        'materials': materials,
    }


def _ensure_template() -> str:
    """模板不存在时运行 tools/build_supply_template.py 生成。"""
    if not os.path.exists(TEMPLATE_PATH):
        script = os.path.join(_ROOT, 'tools', 'build_supply_template.py')
        subprocess.run([sys.executable, script], check=True, cwd=_ROOT)
    return TEMPLATE_PATH


def _iter_all_runs(doc):
    for p in doc.paragraphs:
        yield from p.runs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield from p.runs


def _full_text(doc) -> str:
    return ''.join(run.text for run in _iter_all_runs(doc))


# ---------------------------------------------------------------------------
# 测试
# ---------------------------------------------------------------------------
def test_gen_supply_form():
    """假数据生成供方单：全部占位替换成功、无红色残留、关键文字正确。"""
    template = _ensure_template()
    order, supply = _fake_order(), _fake_supply()

    tmp = tempfile.mkdtemp(prefix='gen_supply_test_')
    try:
        out_path = os.path.join(tmp, 'subdir', '供方单.docx')  # 顺带验证自动建目录
        fmap = _build_field_map(order, supply)
        if fmap is None:
            fmap = _inline_field_map(order, supply)
            unmatched = gen_supply_form(order, supply, template, out_path,
                                        field_map=fmap)
        else:
            unmatched = gen_supply_form(order, supply, template, out_path)

        assert unmatched == [], f'存在未匹配占位: {unmatched}'
        assert os.path.exists(out_path)

        doc = Document(out_path)
        # 不存在任何红色 run
        reds = [r.text for r in _iter_all_runs(doc)
                if r.font.color.rgb is not None
                and str(r.font.color.rgb).upper() == 'FF0000']
        assert reds == [], f'存在未替换的红色 run: {reds}'

        text = _full_text(doc)
        for expect in ('GD20240108001',        # 流程编号
                       '3×220/380',            # 电表电压型号
                       '400A三相一表位落地式表箱',  # 表箱型号
                       ):
            assert expect in text, f'生成结果缺少: {expect}'
        # 施工规格句子（契约/内联两种路径各自的拼句结果都应原样出现）
        assert fmap['施工规格'] in text, f"生成结果缺少施工规格: {fmap['施工规格']}"
    finally:
        shutil.rmtree(tmp, ignore_errors=True)


def test_unmatched_placeholder_reported():
    """field_map 缺键时保留原文并记入返回列表。"""
    template = _ensure_template()
    order, supply = _fake_order(), _fake_supply()
    fmap = _inline_field_map(order, supply)
    del fmap['电价']

    tmp = tempfile.mkdtemp(prefix='gen_supply_test_')
    try:
        out_path = os.path.join(tmp, '供方单.docx')
        unmatched = gen_supply_form(order, supply, template, out_path,
                                    field_map=fmap)
        # 模板中同一占位可能出现多处（合并单元格/重复内容），按占位名集合断言
        assert unmatched and set(unmatched) == {'电价'}, f'未匹配列表不符: {unmatched}'
    finally:
        shutil.rmtree(tmp, ignore_errors=True)


if __name__ == '__main__':
    # pytest 不可用时可直接 python tests/test_gen_supply.py 运行
    funcs = [(name, fn) for name, fn in sorted(globals().items())
             if name.startswith('test_') and callable(fn)]
    failed = 0
    for name, fn in funcs:
        try:
            fn()
            print(f'通过: {name}')
        except AssertionError as e:
            failed += 1
            print(f'失败: {name} -> {e}')
    print(f'\n共 {len(funcs)} 项，失败 {failed} 项')
    sys.exit(1 if failed else 0)
